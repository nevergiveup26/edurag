"""
文档加载器
支持TXT/PDF/Word多种格式文档加载
集成 LangChain 文档加载器，同时保留原生实现作为 fallback
"""
import os
from typing import List, Optional

from core.models import Document
from core.logger import get_logger

logger = get_logger("document_loader")


class DocumentLoader:
    """文档加载器（LangChain + 原生双模式）"""

    # 支持的文件格式
    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".doc", ".md", ".csv", ".json"}

    @classmethod
    def load_file(cls, file_path: str, **kwargs) -> List[Document]:
        """
        加载单个文件（自动选择最佳加载器）

        Args:
            file_path: 文件路径

        Returns:
            Document 列表
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}。支持: {cls.SUPPORTED_EXTENSIONS}")

        # 优先使用 LangChain 加载器
        try:
            docs = cls._load_with_langchain(file_path, ext)
            if docs:
                return docs
        except Exception as e:
            logger.debug(f"LangChain 加载失败，回退到原生加载器: {e}")

        # fallback 到原生加载器
        return cls._load_with_native(file_path, ext)

    @classmethod
    def _load_with_langchain(cls, file_path: str, ext: str) -> List[Document]:
        """使用 LangChain 加载器"""
        from langchain_community.document_loaders import (
            TextLoader,
            PyPDFLoader,
            Docx2txtLoader,
            UnstructuredMarkdownLoader,
            CSVLoader,
            JSONLoader,
        )

        loader_map = {
            ".txt": TextLoader,
            ".pdf": PyPDFLoader,
            ".docx": Docx2txtLoader,
            ".doc": Docx2txtLoader,
            ".md": UnstructuredMarkdownLoader,
            ".csv": CSVLoader,
            ".json": JSONLoader,
        }

        loader_cls = loader_map.get(ext)
        if loader_cls is None:
            return []

        loader = loader_cls(file_path, encoding="utf-8")
        langchain_docs = loader.load()
        logger.info(f"[LangChain] 加载 {file_path}: {len(langchain_docs)} 个文档")

        docs = []
        for i, lc_doc in enumerate(langchain_docs):
            doc_id = f"{os.path.basename(file_path)}_lc_{i}"
            metadata = dict(lc_doc.metadata) if lc_doc.metadata else {}
            metadata["loader"] = "langchain"
            docs.append(Document(
                content=lc_doc.page_content,
                doc_id=doc_id,
                title=os.path.basename(file_path),
                source=file_path,
                metadata=metadata,
            ))
        return docs

    @classmethod
    def _load_with_native(cls, file_path: str, ext: str) -> List[Document]:
        """使用原生加载器（fallback）"""
        logger.info(f"[原生] 加载 {file_path}")

        if ext == ".txt" or ext == ".md":
            return cls._load_txt(file_path)
        elif ext == ".pdf":
            return cls._load_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return cls._load_word(file_path)
        elif ext == ".csv":
            return cls._load_csv(file_path)
        elif ext == ".json":
            return cls._load_json(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    @classmethod
    def _load_txt(cls, file_path: str, encoding: str = "utf-8") -> List[Document]:
        """加载TXT/MD文件"""
        with open(file_path, "r", encoding=encoding) as f:
            content = f.read()
        return [Document(
            content=content,
            doc_id=os.path.basename(file_path),
            title=os.path.basename(file_path),
            source=file_path,
        )]

    @classmethod
    def _load_pdf(cls, file_path: str) -> List[Document]:
        """加载PDF文件"""
        import PyPDF2
        docs = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(Document(
                        content=text,
                        doc_id=f"{os.path.basename(file_path)}_p{i+1}",
                        title=os.path.basename(file_path),
                        source=file_path,
                        metadata={"page": i + 1},
                    ))
        logger.info(f"加载PDF: {file_path}, {len(docs)}页")
        return docs

    @classmethod
    def _load_word(cls, file_path: str) -> List[Document]:
        """加载Word文件"""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n".join(paragraphs)
        return [Document(
            content=content,
            doc_id=os.path.basename(file_path),
            title=os.path.basename(file_path),
            source=file_path,
        )]

    @classmethod
    def _load_csv(cls, file_path: str) -> List[Document]:
        """加载CSV文件"""
        import csv
        docs = []
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for i, row in enumerate(reader):
                content = "\n".join([f"{k}: {v}" for k, v in row.items()])
                docs.append(Document(
                    content=content,
                    doc_id=f"{os.path.basename(file_path)}_row{i}",
                    source=file_path,
                    metadata={"row": i},
                ))
        return docs

    @classmethod
    def _load_json(cls, file_path: str) -> List[Document]:
        """加载JSON文件"""
        import json
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        content = json.dumps(data, ensure_ascii=False, indent=2)
        return [Document(
            content=content,
            doc_id=os.path.basename(file_path),
            source=file_path,
        )]

    @classmethod
    def load_directory(cls, directory: str, extensions: List[str] = None,
                       recursive: bool = True) -> List[Document]:
        """加载目录下所有支持的文档"""
        if extensions is None:
            extensions = list(cls.SUPPORTED_EXTENSIONS)

        all_docs = []
        walker = os.walk(directory) if recursive else [(directory, [], os.listdir(directory))]

        for root, _, files in walker:
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in extensions:
                    file_path = os.path.join(root, file)
                    try:
                        docs = cls.load_file(file_path)
                        all_docs.extend(docs)
                        logger.info(f"已加载: {file_path}")
                    except Exception as e:
                        logger.error(f"加载失败 {file_path}: {e}")

        logger.info(f"目录加载完成: {len(all_docs)} 个文档")
        return all_docs