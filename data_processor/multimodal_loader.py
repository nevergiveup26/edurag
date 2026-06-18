"""
多模态文档加载器

从PDF、DOCX、PPTX等文档中提取：
- 纯文本内容
- 图片（保存到本地 + 生成base64）
- 表格（提取为结构化JSON）
- 图表标题与描述

支持的后端：
- PyMuPDF (fitz) — PDF图片提取
- python-docx — DOCX图片提取
- langchain_unstructured — 通用文档解析

使用方式：
    加载文档 → 切分 → 生成MultiModalChunk → 存入向量库
"""
import os
import base64
import json
import uuid
from typing import List, Optional, Dict, Any, Tuple
from pathlib import Path

from core.logger import get_logger
from core.models import MultiModalChunk, ContentType

logger = get_logger("multimodal_loader")

# 临时图片存储目录
MEDIA_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "media")


class MultiModalLoader:
    """多模态文档加载器"""

    def __init__(self, media_dir: str = None):
        self.media_dir = media_dir or MEDIA_DIR
        os.makedirs(self.media_dir, exist_ok=True)

    # ======================== 文本提取 ========================

    def extract_text(self, file_path: str) -> str:
        """
        从文档中提取纯文本

        支持: .txt, .pdf, .docx, .md, .html
        """
        ext = Path(file_path).suffix.lower()

        if ext == ".txt" or ext == ".md":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        if ext == ".pdf":
            return self._extract_pdf_text(file_path)

        if ext == ".docx":
            return self._extract_docx_text(file_path)

        if ext == ".html":
            return self._extract_html_text(file_path)

        logger.warning(f"不支持的文件格式: {ext}")
        return ""

    # ======================== 图片提取 ========================

    def extract_images(self, file_path: str) -> List[Dict[str, Any]]:
        """
        从文档中提取所有图片

        Returns:
            [{"image_path": "...", "image_base64": "...", "page": 1, "caption": "..."}, ...]
        """
        ext = Path(file_path).suffix.lower()
        images = []

        if ext == ".pdf":
            images = self._extract_pdf_images(file_path)
        elif ext == ".docx":
            images = self._extract_docx_images(file_path)
        else:
            logger.debug(f"暂不支持 {ext} 图片提取")

        return images

    # ======================== 表格提取 ========================

    def extract_tables(self, file_path: str) -> List[Dict[str, Any]]:
        """
        从文档中提取所有表格

        Returns:
            [{"table_data": [[...], ...], "page": 2, "caption": "..."}, ...]
        """
        ext = Path(file_path).suffix.lower()
        tables = []

        if ext == ".pdf":
            tables = self._extract_pdf_tables(file_path)
        elif ext == ".docx":
            tables = self._extract_docx_tables(file_path)
        else:
            logger.debug(f"暂不支持 {ext} 表格提取")

        return tables

    # ======================== 完整多模态加载 ========================

    def load_document(self, file_path: str, chunk_size: int = 500) -> List[MultiModalChunk]:
        """
        完整加载文档，生成多模态片段列表

        流程：
        1. 加载文本 → 切分为文本块
        2. 提取图片 → 生成图片块（含描述占位）
        3. 提取表格 → 生成表格块
        4. 按文档出现顺序组合

        Returns:
            MultiModalChunk列表（按在文档中的顺序排列）
        """
        chunks = []
        doc_id = str(uuid.uuid4())
        file_name = Path(file_path).name

        # 1. 文本块
        text = self.extract_text(file_path)
        if text:
            text_chunks = self._split_text(text, chunk_size)
            for i, tc in enumerate(text_chunks):
                chunks.append(MultiModalChunk(
                    chunk_id=f"{doc_id}_text_{i}",
                    content=tc,
                    content_type=ContentType.TEXT.value,
                    doc_id=doc_id,
                    metadata={"source": file_name, "page": 0, "type": "text"},
                ))

        # 2. 图片块
        images = self.extract_images(file_path)
        for i, img in enumerate(images):
            chunks.append(MultiModalChunk(
                chunk_id=f"{doc_id}_img_{i}",
                content=img.get("caption", f"图片 {i+1}"),
                content_type=ContentType.IMAGE.value,
                doc_id=doc_id,
                image_path=img.get("image_path"),
                image_base64=img.get("image_base64"),
                image_description=img.get("caption", ""),
                metadata={
                    "source": file_name,
                    "page": img.get("page", 0),
                    "type": "image",
                    "image_index": i,
                },
            ))

        # 3. 表格块
        tables = self.extract_tables(file_path)
        for i, tbl in enumerate(tables):
            table_str = json.dumps(tbl.get("table_data", []), ensure_ascii=False)
            chunks.append(MultiModalChunk(
                chunk_id=f"{doc_id}_tbl_{i}",
                content=table_str,
                content_type=ContentType.TABLE.value,
                doc_id=doc_id,
                table_data=table_str,
                metadata={
                    "source": file_name,
                    "page": tbl.get("page", 0),
                    "type": "table",
                    "caption": tbl.get("caption", ""),
                },
            ))

        logger.info(f"多模态加载完成: {file_name} → {len(chunks)} 块 "
                     f"(文本:{len([c for c in chunks if c.content_type=='text'])}, "
                     f"图片:{len([c for c in chunks if c.content_type=='image'])}, "
                     f"表格:{len([c for c in chunks if c.content_type=='table'])})")
        return chunks

    # ======================== PDF处理 ========================

    def _extract_pdf_text(self, file_path: str) -> str:
        """从PDF提取文本"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            logger.warning("PyMuPDF未安装，尝试使用 langchain")
        try:
            # LangChain fallback
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(file_path)
            pages = loader.load()
            return "\n".join(p.page_content for p in pages)
        except Exception as e:
            logger.error(f"PDF文本提取失败: {e}")
            return ""

    def _extract_pdf_images(self, file_path: str) -> List[Dict]:
        """从PDF提取图片"""
        images = []
        try:
            import fitz
            doc = fitz.open(file_path)
            for page_idx, page in enumerate(doc):
                image_list = page.get_images(full=True)
                for img_idx, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    ext = base_image["ext"]

                    # 保存到本地
                    img_filename = f"{Path(file_path).stem}_p{page_idx+1}_img{img_idx}.{ext}"
                    img_path = os.path.join(self.media_dir, img_filename)
                    with open(img_path, "wb") as f:
                        f.write(image_bytes)

                    images.append({
                        "image_path": img_path,
                        "image_base64": base64.b64encode(image_bytes).decode(),
                        "page": page_idx + 1,
                        "width": base_image.get("width"),
                        "height": base_image.get("height"),
                        "caption": f"图片 (第{page_idx+1}页)",
                    })
            doc.close()
            logger.info(f"PDF图片提取: {len(images)}张")
        except Exception as e:
            logger.warning(f"PDF图片提取失败: {e}")
        return images

    def _extract_pdf_tables(self, file_path: str) -> List[Dict]:
        """从PDF提取表格"""
        tables = []
        try:
            import fitz
            doc = fitz.open(file_path)
            for page_idx, page in enumerate(doc):
                # PyMuPDF 的表格检测
                tabs = page.find_tables()
                if tabs and tabs.tables:
                    for tab_idx, tab in enumerate(tabs.tables):
                        rows = []
                        for row in tab.extract():
                            rows.append([str(cell) if cell is not None else "" for cell in row])
                        tables.append({
                            "table_data": rows,
                            "page": page_idx + 1,
                            "caption": f"表格 (第{page_idx+1}页)",
                            "rows": len(rows),
                            "cols": len(rows[0]) if rows else 0,
                        })
            doc.close()
            logger.info(f"PDF表格提取: {len(tables)}个")
        except ImportError:
            logger.debug("PyMuPDF未安装，跳过表格提取")
        except Exception as e:
            logger.warning(f"PDF表格提取失败: {e}")
        return tables

    # ======================== DOCX处理 ========================

    def _extract_docx_text(self, file_path: str) -> str:
        """从DOCX提取文本"""
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            logger.warning("python-docx未安装")
        except Exception as e:
            logger.error(f"DOCX文本提取失败: {e}")
        return ""

    def _extract_docx_images(self, file_path: str) -> List[Dict]:
        """从DOCX提取图片"""
        images = []
        try:
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            doc = Document(file_path)
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image = rel.target_part
                    image_bytes = image.blob
                    ext = image.content_type.split("/")[-1] if "/" in image.content_type else "png"

                    img_filename = f"{Path(file_path).stem}_{uuid.uuid4().hex[:8]}.{ext}"
                    img_path = os.path.join(self.media_dir, img_filename)
                    with open(img_path, "wb") as f:
                        f.write(image_bytes)

                    images.append({
                        "image_path": img_path,
                        "image_base64": base64.b64encode(image_bytes).decode(),
                        "page": 0,
                        "caption": f"图片 ({Path(file_path).stem})",
                    })
            logger.info(f"DOCX图片提取: {len(images)}张")
        except Exception as e:
            logger.warning(f"DOCX图片提取失败: {e}")
        return images

    def _extract_docx_tables(self, file_path: str) -> List[Dict]:
        """从DOCX提取表格"""
        tables = []
        try:
            from docx import Document
            doc = Document(file_path)
            for tab_idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                tables.append({
                    "table_data": rows,
                    "page": 0,
                    "caption": f"表格 {tab_idx+1}",
                    "rows": len(rows),
                    "cols": len(rows[0]) if rows else 0,
                })
            logger.info(f"DOCX表格提取: {len(tables)}个")
        except Exception as e:
            logger.warning(f"DOCX表格提取失败: {e}")
        return tables

    # ======================== HTML处理 ========================

    def _extract_html_text(self, file_path: str) -> str:
        try:
            from bs4 import BeautifulSoup
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            logger.warning("beautifulsoup4未安装")
        except Exception as e:
            logger.error(f"HTML文本提取失败: {e}")
        return ""

    # ======================== 文本切分 ========================

    def _split_text(self, text: str, chunk_size: int = 500) -> List[str]:
        """简单段落切分"""
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
        chunks = []
        current = ""
        for para in paragraphs:
            if len(current) + len(para) < chunk_size:
                current += para + "\n"
            else:
                if current:
                    chunks.append(current.strip())
                current = para + "\n"
        if current.strip():
            chunks.append(current.strip())
        return chunks if chunks else [text[:chunk_size]]


# ======================== 直接加载图片文件 ========================

    def load_image(self, file_path: str, caption: str = "") -> MultiModalChunk:
        """
        直接加载单张图片文件（.jpg/.png/.gif/.webp/.bmp）

        与 load_document() 不同，此方法不提取文本，而是将整个图片作为
        一个 IMAGE 类型的 MultiModalChunk 返回。

        Returns:
            单个 MultiModalChunk (content_type=IMAGE)
        """
        file_name = Path(file_path).name
        doc_id = str(uuid.uuid4())

        with open(file_path, "rb") as f:
            image_bytes = f.read()

        img_base64 = base64.b64encode(image_bytes).decode()

        # 尝试获取图片宽高
        width, height = 0, 0
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                width, height = img.size
        except Exception as e:
            logger.debug(f"图片尺寸获取失败: {e}")

        chunk = MultiModalChunk(
            chunk_id=f"{doc_id}_img_0",
            content=caption or f"图片: {file_name}",
            content_type=ContentType.IMAGE.value,
            doc_id=doc_id,
            image_path=file_path,
            image_base64=img_base64,
            image_description=caption or file_name,
            metadata={
                "source": file_name,
                "type": "image",
                "width": width,
                "height": height,
                "caption": caption or file_name,
            },
        )

        logger.info(f"图片加载完成: {file_name} ({width}x{height}, {len(image_bytes)} bytes)")
        return chunk


def extract_multimodal_chunks(file_path: str, chunk_size: int = 500) -> List[MultiModalChunk]:
    """便捷函数：提取多模态文档块"""
    loader = MultiModalLoader()
    return loader.load_document(file_path, chunk_size)